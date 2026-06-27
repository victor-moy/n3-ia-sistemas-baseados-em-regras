"""
Converte relatorio.md para relatorio.docx com formatação ABNT básica.
"""

import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_font(run, name="Times New Roman", size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_paragraph_format(para, space_before=0, space_after=6, line_spacing=1.5, alignment=None):
    fmt = para.paragraph_format
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)
    fmt.line_spacing = Pt(12 * line_spacing)
    if alignment is not None:
        fmt.alignment = alignment


def add_heading(doc, text, level):
    colors = {1: (0, 0, 0), 2: (0, 0, 0), 3: (50, 50, 50)}
    sizes = {1: 14, 2: 13, 3: 12}

    para = doc.add_paragraph()
    run = para.add_run(text)
    set_font(run, size=sizes.get(level, 12), bold=True, color=colors.get(level, (0, 0, 0)))
    fmt = para.paragraph_format
    fmt.space_before = Pt(18 if level == 1 else 12)
    fmt.space_after = Pt(6)
    fmt.line_spacing = Pt(18)
    fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return para


def add_body(doc, text, bold_parts=None, italic_parts=None):
    para = doc.add_paragraph()
    set_paragraph_format(para, space_before=0, space_after=6, line_spacing=1.5,
                         alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # processa negrito/itálico inline com markdown
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            set_font(run, bold=True)
        elif part.startswith("*") and part.endswith("*"):
            run = para.add_run(part[1:-1])
            set_font(run, italic=True)
        elif part.startswith("`") and part.endswith("`"):
            run = para.add_run(part[1:-1])
            set_font(run, name="Courier New", size=10)
        else:
            run = para.add_run(part)
            set_font(run)
    return para


def add_code(doc, lines):
    para = doc.add_paragraph()
    fmt = para.paragraph_format
    fmt.space_before = Pt(6)
    fmt.space_after = Pt(6)
    fmt.left_indent = Cm(1)

    # fundo cinza via shading XML
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)

    code_text = "\n".join(lines)
    run = para.add_run(code_text)
    set_font(run, name="Courier New", size=9)
    return para


def add_table(doc, rows):
    if not rows:
        return
    # detecta separador (linha com ---)
    data = [r for r in rows if not re.match(r'^\|[-| :]+\|$', r)]
    if len(data) < 2:
        return

    def parse_row(row):
        return [c.strip() for c in row.strip().strip('|').split('|')]

    headers = parse_row(data[0])
    body = [parse_row(r) for r in data[1:]]

    cols = len(headers)
    table = doc.add_table(rows=1 + len(body), cols=cols)
    table.style = 'Table Grid'

    # cabeçalho
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(h)
        if cell.paragraphs[0].runs:
            set_font(cell.paragraphs[0].runs[0], bold=True, size=10)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'DDDDDD')
        cell._tc.get_or_add_tcPr().append(shd)

    # dados
    for ri, row_data in enumerate(body):
        trow = table.rows[ri + 1]
        for ci, cell_text in enumerate(row_data[:cols]):
            cell = trow.cells[ci]
            cell.text = cell_text
            if cell.paragraphs[0].runs:
                set_font(cell.paragraphs[0].runs[0], size=10)

    doc.add_paragraph()  # espaço após tabela


def add_bullet(doc, text, level=0):
    para = doc.add_paragraph(style='List Bullet')
    fmt = para.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(3)
    fmt.line_spacing = Pt(18)
    fmt.left_indent = Cm(1 + level * 0.5)

    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            set_font(run, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = para.add_run(part[1:-1])
            set_font(run, name="Courier New", size=10)
        else:
            run = para.add_run(part)
            set_font(run)
    return para


def add_cover(doc):
    # Instituição
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SISTEMAS BASEADOS EM REGRAS")
    set_font(r, size=16, bold=True)
    p.paragraph_format.space_before = Pt(72)
    p.paragraph_format.space_after = Pt(6)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Aplicações Inteligentes no Dia a Dia")
    set_font(r2, size=14)
    p2.paragraph_format.space_after = Pt(60)

    for line in [
        "Disciplina: Inteligência Artificial",
        "Professor: Claudinei Dias (Ney)",
        "Semestre: 2026/1",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        set_font(r, size=12)
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()


def parse_and_build(md_path, docx_path):
    with open(md_path, encoding="utf-8") as f:
        lines = f.readlines()

    doc = Document()

    # Margens ABNT: sup/inf 2,5cm, esq 3cm, dir 2cm
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)

    add_cover(doc)

    i = 0
    code_buffer = []
    table_buffer = []
    in_code = False
    in_table = False

    # pula o bloco de cabeçalho YAML/markdown (título + metadados)
    skip_header = True

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip("\n")

        # pula bloco inicial de cabeçalho até o primeiro ---
        if skip_header:
            if line.strip() == "---" and i > 0:
                skip_header = False
            i += 1
            continue

        # bloco de código
        if line.startswith("```"):
            if not in_code:
                in_code = True
                code_buffer = []
            else:
                in_code = False
                add_code(doc, code_buffer)
                code_buffer = []
            i += 1
            continue

        if in_code:
            code_buffer.append(line)
            i += 1
            continue

        # tabela markdown
        if line.startswith("|"):
            if not in_table:
                in_table = True
                table_buffer = []
            table_buffer.append(line)
            i += 1
            continue
        else:
            if in_table:
                in_table = False
                add_table(doc, table_buffer)
                table_buffer = []

        # separador ---
        if re.match(r'^---+$', line.strip()):
            i += 1
            continue

        # headings
        m = re.match(r'^(#{1,3})\s+(.*)', line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            # limpa markdown inline do heading
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            text = re.sub(r'\*([^*]+)\*', r'\1', text)
            add_heading(doc, text, level)
            i += 1
            continue

        # lista com bullet
        m_bullet = re.match(r'^(\s*)[-*]\s+(.*)', line)
        if m_bullet:
            indent = len(m_bullet.group(1)) // 2
            add_bullet(doc, m_bullet.group(2).strip(), level=indent)
            i += 1
            continue

        # lista numerada
        m_num = re.match(r'^\d+\.\s+(.*)', line)
        if m_num:
            add_bullet(doc, m_num.group(1).strip())
            i += 1
            continue

        # linha em branco
        if line.strip() == "":
            i += 1
            continue

        # parágrafo normal
        add_body(doc, line.strip())
        i += 1

    # flush de table pendente
    if in_table and table_buffer:
        add_table(doc, table_buffer)

    doc.save(docx_path)
    print(f"Arquivo salvo: {docx_path}")


if __name__ == "__main__":
    import os
    base = os.path.dirname(os.path.abspath(__file__))
    parse_and_build(
        os.path.join(base, "relatorio.md"),
        os.path.join(base, "relatorio.docx"),
    )
